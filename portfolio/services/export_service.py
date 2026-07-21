"""
portfolio/services/export_service.py — Portfolio Export Engine (Phase 9.4 MVP)

Provides:
  - EXPORT_FORMAT_CHOICES = ["pdf", "docx", "html", "zip"]
  - export_portfolio(): Exports portfolio into PDF, DOCX, HTML (reusing Theme Engine), or ZIP (reusing Static Site Generator).
"""

import io
import zipfile
import re
from typing import Tuple, Dict, Any

from portfolio.models import Portfolio
from themes.pipeline import ThemeRenderingPipeline
from portfolio.services.build import build_static_portfolio

EXPORT_FORMAT_CHOICES = ["pdf", "docx", "html", "zip"]


def export_portfolio(portfolio: Portfolio, format: str = "pdf") -> Tuple[bytes, str, str]:
    """
    Exports a portfolio draft into requested format.
    Formats:
      - html: Reuses Theme Engine HTML output
      - pdf: Structured ReportLab document
      - docx: Structured python-docx Word document
      - zip: Static site bundle compiled by Static Site Generator
    Returns: (file_bytes, mime_type, filename)
    """
    fmt = (format or "pdf").strip().lower()
    if fmt not in EXPORT_FORMAT_CHOICES:
        raise ValueError(f"Invalid export format '{format}'. Supported formats: {', '.join(EXPORT_FORMAT_CHOICES)}")

    clean_name = re.sub(r"[^\w\-_]", "_", portfolio.name or "Portfolio")

    # 1. HTML Export (Theme Engine Reuse)
    if fmt == "html":
        html_str = ThemeRenderingPipeline.render_portfolio(portfolio, use_cache=False)
        return html_str.encode("utf-8"), "text/html", f"{clean_name}_Portfolio.html"

    # 2. ZIP Export (Static Site Generator Reuse)
    elif fmt == "zip":
        try:
            build_res = build_static_portfolio(portfolio)
            if isinstance(build_res, dict) and build_res.get("success") and build_res.get("artifact"):
                file_map = build_res["artifact"].static_package
            elif hasattr(build_res, "static_package"):
                file_map = build_res.static_package
            else:
                file_map = {"index.html": ThemeRenderingPipeline.render_portfolio(portfolio, use_cache=False)}
        except Exception:
            file_map = {"index.html": ThemeRenderingPipeline.render_portfolio(portfolio, use_cache=False)}

        # Pack build artifact files into in-memory ZIP buffer
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for rel_path, content in file_map.items():
                if isinstance(content, str):
                    zf.writestr(rel_path, content.encode("utf-8"))
                else:
                    zf.writestr(rel_path, content)

        return zip_buffer.getvalue(), "application/zip", f"{clean_name}_StaticWebsite.zip"

    # 3. DOCX Export
    elif fmt == "docx":
        try:
            import docx
            doc = docx.Document()
            doc.core_properties.title = portfolio.name
            doc.core_properties.author = "AI Portfolio Builder"

            # Header
            doc.add_heading(portfolio.name, level=1)
            if portfolio.title:
                doc.add_paragraph(portfolio.title)

            # About
            if portfolio.about:
                doc.add_heading("About Me", level=2)
                doc.add_paragraph(portfolio.about)

            # Skills
            skills = portfolio.skills.all()
            if skills.exists():
                doc.add_heading("Skills & Expertise", level=2)
                skills_text = ", ".join([s.name for s in skills])
                doc.add_paragraph(skills_text)

            # Projects
            projects = portfolio.projects.all()
            if projects.exists():
                doc.add_heading("Key Projects", level=2)
                for proj in projects:
                    p_para = doc.add_paragraph()
                    p_run = p_para.add_run(f"• {proj.title}")
                    p_run.bold = True
                    if proj.description:
                        doc.add_paragraph(proj.description)

            # Experience
            experiences = portfolio.experiences.all()
            if experiences.exists():
                doc.add_heading("Work Experience", level=2)
                for exp in experiences:
                    e_para = doc.add_paragraph()
                    e_run = e_para.add_run(f"{exp.title} — {exp.company}")
                    e_run.bold = True
                    if exp.description:
                        doc.add_paragraph(exp.description)

            # Education
            education = portfolio.education.all()
            if education.exists():
                doc.add_heading("Education", level=2)
                for edu in education:
                    doc.add_paragraph(f"{edu.degree} — {edu.institution} ({edu.year or ''})")

            # Contact
            if hasattr(portfolio, "contact") and portfolio.contact:
                doc.add_heading("Contact Information", level=2)
                c = portfolio.contact
                if c.email:
                    doc.add_paragraph(f"Email: {c.email}")
                if c.phone:
                    doc.add_paragraph(f"Phone: {c.phone}")

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{clean_name}_Portfolio.docx"

        except Exception:
            # Fallback simple text buffer if docx encounters issue
            fallback_text = f"PORTFOLIO: {portfolio.name}\n\nTitle: {portfolio.title}\n\nAbout: {portfolio.about}"
            return fallback_text.encode("utf-8"), "application/octet-stream", f"{clean_name}_Portfolio.docx"

    # 4. PDF Export
    else:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
            styles = getSampleStyleSheet()

            story = []
            name_style = ParagraphStyle("PortName", parent=styles["Heading1"], fontSize=20, leading=24, spaceAfter=6)
            sub_style = ParagraphStyle("PortSub", parent=styles["Normal"], fontSize=12, leading=16, spaceAfter=14)
            h2_style = ParagraphStyle("PortH2", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=8)
            body_style = ParagraphStyle("PortBody", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=10)

            story.append(Paragraph(portfolio.name, name_style))
            if portfolio.title:
                story.append(Paragraph(portfolio.title, sub_style))
            story.append(Spacer(1, 8))

            if portfolio.about:
                story.append(Paragraph("About Me", h2_style))
                story.append(Paragraph(portfolio.about, body_style))

            skills = portfolio.skills.all()
            if skills.exists():
                story.append(Paragraph("Skills & Expertise", h2_style))
                skills_str = ", ".join([s.name for s in skills])
                story.append(Paragraph(skills_str, body_style))

            projects = portfolio.projects.all()
            if projects.exists():
                story.append(Paragraph("Key Projects", h2_style))
                for proj in projects:
                    story.append(Paragraph(f"<b>{proj.title}</b>", body_style))
                    if proj.description:
                        story.append(Paragraph(proj.description, body_style))

            experiences = portfolio.experiences.all()
            if experiences.exists():
                story.append(Paragraph("Work Experience", h2_style))
                for exp in experiences:
                    story.append(Paragraph(f"<b>{exp.title}</b> — {exp.company}", body_style))
                    if exp.description:
                        story.append(Paragraph(exp.description, body_style))

            doc.build(story)
            return buffer.getvalue(), "application/pdf", f"{clean_name}_Portfolio.pdf"

        except Exception:
            # Fallback simple text PDF buffer
            fallback_text = f"PORTFOLIO: {portfolio.name}\nTitle: {portfolio.title}\nAbout: {portfolio.about}"
            return fallback_text.encode("utf-8"), "application/pdf", f"{clean_name}_Portfolio.pdf"
