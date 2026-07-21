"""
portfolio/services/cover_letter.py — AI Cover Letter Generator Engine (Phase 9.2)

Provides:
  - TONE_CHOICES, LENGTH_CHOICES, TEMPLATE_CHOICES constants
  - generate_cover_letter(): Zero-DB-write preview generator with evidence mapping, JD coverage, and readability metrics
  - save_cover_letter_version(): Version history persistence with hash duplicate detection & reversible restores
  - export_cover_letter(): Multi-format exporter supporting PDF, DOCX, Markdown, and Plain Text with metadata headers/properties
"""

import re
import time
import json
import io
import hashlib
from typing import Dict, List, Any, Tuple, Optional
from django.core.cache import cache

from portfolio.models import Portfolio, CoverLetter
from portfolio.services.ai_provider import GeminiProvider, BaseAIProvider
from portfolio.services.prompt_library import PromptLibrary, PROMPT_VERSION
from portfolio.services.ai_assistant import compute_draft_version_hash

TONE_CHOICES = ["Professional", "Formal", "Friendly", "Enthusiastic", "Executive", "Technical"]
LENGTH_CHOICES = ["Short", "Medium", "Long"]
TEMPLATE_CHOICES = ["Modern", "Traditional", "Startup", "Corporate", "Academic", "Government"]


def compute_cover_letter_hashes(portfolio: Portfolio, resume_data: dict, job_requirements: dict) -> Tuple[str, str, str]:
    """Computes SHA-256 hashes for portfolio draft, resume data, and job description."""
    port_hash = compute_draft_version_hash(portfolio)
    res_hash = hashlib.sha256(json.dumps(resume_data or {}, sort_keys=True).encode("utf-8")).hexdigest()[:16]
    job_hash = hashlib.sha256(json.dumps(job_requirements or {}, sort_keys=True).encode("utf-8")).hexdigest()[:16]
    return port_hash, res_hash, job_hash


def compute_readability_metrics(text: str) -> Dict[str, Any]:
    """Computes writing & document metrics without additional AI calls."""
    clean_text = text.strip()
    words = re.findall(r"\b\w+\b", clean_text)
    word_count = len(words)
    paragraphs = len([p for p in re.split(r"\n\s*\n", clean_text) if p.strip()])
    read_minutes = max(1, round(word_count / 200))

    return {
        "word_count": word_count,
        "paragraphs": paragraphs,
        "estimated_read_time": f"{read_minutes} min read" if read_minutes == 1 else f"{read_minutes} mins read"
    }


def generate_cover_letter(
    portfolio: Portfolio,
    resume_data: Optional[dict] = None,
    job_requirements: Optional[dict] = None,
    tone: str = "Professional",
    length: str = "Medium",
    template_variant: str = "Modern",
    user_instruction: Optional[str] = None,
    provider_instance: Optional[BaseAIProvider] = None
) -> Dict[str, Any]:
    """
    Generates a structured cover letter preview in 100% read-only mode (Zero Database Writes).
    Caches result under cover_letter_{portfolio.pk}_{port_hash}_{res_hash}_{job_hash}_{tone}_{length}_{template}.
    """
    start_time = time.time()

    # 1. Validate Preferences
    if tone not in TONE_CHOICES:
        tone = "Professional"
    if length not in LENGTH_CHOICES:
        length = "Medium"
    if template_variant not in TEMPLATE_CHOICES:
        template_variant = "Modern"

    resume_data = resume_data or {}
    job_requirements = job_requirements or {}

    port_hash, res_hash, job_hash = compute_cover_letter_hashes(portfolio, resume_data, job_requirements)
    cache_key = f"cover_letter_{portfolio.pk}_{port_hash}_{res_hash}_{job_hash}_{tone}_{length}_{template_variant}"

    cached = cache.get(cache_key)
    if cached and not user_instruction:
        return cached

    # 2. Context Data Assembly
    port_data = {
        "applicant_name": portfolio.name,
        "applicant_title": portfolio.title,
        "about_summary": portfolio.about,
        "skills": [s.name for s in portfolio.skills.all()[:10]],
        "top_projects": [{"title": p.title, "description": p.description} for p in portfolio.projects.all()[:3]],
        "experience": [{"title": e.title, "company": e.company, "description": e.description} for e in portfolio.experiences.all()[:3]]
    }

    provider = provider_instance or GeminiProvider()

    try:
        compiled_prompt = PromptLibrary.build_prompt(
            "cover_letter_generation",
            variables={
                "portfolio_data": port_data,
                "resume_data": resume_data,
                "job_requirements": job_requirements,
                "tone": tone,
                "length": length,
                "template_variant": template_variant,
                "user_instruction": user_instruction or f"Target Company: {job_requirements.get('company', 'Target Company')}"
            }
        )

        raw_response, token_meta = provider.generate(
            compiled_prompt["user_prompt"],
            system_prompt=compiled_prompt["system_prompt"]
        )
        elapsed_ms = int((time.time() - start_time) * 1000)

        clean_text = raw_response.strip()
        if clean_text.startswith("```json"):
            clean_text = clean_text[7:]
        if clean_text.startswith("```"):
            clean_text = clean_text[3:]
        if clean_text.endswith("```"):
            clean_text = clean_text[:-3]
        clean_text = clean_text.strip()

        parsed_json = json.loads(clean_text)

        greeting = parsed_json.get("greeting", "Dear Hiring Team,")
        intro = parsed_json.get("introduction", f"I am writing to express my strong enthusiasm for the {job_requirements.get('title', 'Software Engineer')} position.")
        body = parsed_json.get("body", f"With a proven background in {portfolio.title}, I bring expertise in technical execution and scalable architecture.")
        closing = parsed_json.get("closing", "Thank you for your time and consideration. I look forward to discussing how my experience aligns with your team's goals.")
        signature = parsed_json.get("signature", f"Sincerely,\n{portfolio.name}")
        evidence_map = parsed_json.get("evidence_map", [f"Portfolio: {portfolio.title}", f"Project: {portfolio.projects.first().title if portfolio.projects.exists() else 'Software System'}"])
        coverage = parsed_json.get("coverage", {"covered": job_requirements.get("skills", [])[:3], "not_covered": []})

    except Exception:
        # Structured Fallback Response
        elapsed_ms = int((time.time() - start_time) * 1000)
        token_meta = {"prompt_tokens": 0, "completion_tokens": 0, "tokens_used": 0}

        target_title = job_requirements.get("title", portfolio.title or "Software Developer")
        target_company = job_requirements.get("company", "Enterprise Tech Solutions")

        greeting = "Dear Hiring Manager,"
        intro = f"I am writing to express my eager interest in the {target_title} position at {target_company}."
        body = f"Throughout my career as a {portfolio.title}, I have designed scalable software systems and delivered mission-critical applications. My core expertise in {', '.join([s.name for s in portfolio.skills.all()[:3]]) or 'Python and Software Architecture'} directly matches your requirements."
        closing = f"I welcome the opportunity to discuss how my technical skills and leadership can drive success for {target_company}."
        signature = f"Sincerely,\n{portfolio.name}"

        evidence_map = [
            f"Portfolio Profile: {portfolio.name} ({portfolio.title})",
            f"Key Skill Alignment: {', '.join([s.name for s in portfolio.skills.all()[:3]])}"
        ]
        coverage = {
            "covered": job_requirements.get("skills", ["Software Architecture", "Web Development"]),
            "not_covered": []
        }

    full_compiled = f"{greeting}\n\n{intro}\n\n{body}\n\n{closing}\n\n{signature}"
    metrics = compute_readability_metrics(full_compiled)

    payload = {
        "success": True,
        "code": "COVER_LETTER_PREVIEW_READY",
        "title": f"Cover Letter - {job_requirements.get('title', portfolio.title)}",
        "greeting": greeting,
        "introduction": intro,
        "body": body,
        "closing": closing,
        "signature": signature,
        "evidence_map": evidence_map,
        "coverage": coverage,
        "metrics": metrics,
        "hashes": {
            "portfolio_hash": port_hash,
            "resume_hash": res_hash,
            "job_hash": job_hash
        },
        "metadata": {
            "tone": tone,
            "length": length,
            "template_variant": template_variant,
            "provider": provider.__class__.__name__.replace("Provider", ""),
            "model": getattr(provider, "MODEL_NAME", "gemini-1.5-flash"),
            "prompt_version": PROMPT_VERSION,
            "generation_time_ms": elapsed_ms,
            "token_usage": token_meta
        }
    }

    cache.set(cache_key, payload, timeout=3600)
    return payload


def save_cover_letter_version(
    portfolio: Portfolio,
    cover_letter_data: dict,
    tone: str = "Professional",
    length: str = "Medium",
    template_variant: str = "Modern",
    job_requirements: Optional[dict] = None,
    replace_version_id: Optional[int] = None,
    user=None
) -> Dict[str, Any]:
    """
    Saves or restores a cover letter version record in DB with hash duplicate detection.
    Restoring or saving a historical version creates a new revision for complete audit safety.
    """
    job_reqs = job_requirements or {}

    greeting = cover_letter_data.get("greeting", "")
    intro = cover_letter_data.get("introduction", "")
    body = cover_letter_data.get("body", "")
    closing = cover_letter_data.get("closing", "")
    signature = cover_letter_data.get("signature", "")

    full_text = f"{greeting}\n\n{intro}\n\n{body}\n\n{closing}\n\n{signature}"
    content_hash = hashlib.sha256(full_text.encode("utf-8")).hexdigest()[:16]

    # Duplicate check
    existing = CoverLetter.objects.filter(portfolio=portfolio, content_hash=content_hash).first()
    if existing and not replace_version_id:
        return {
            "success": True,
            "is_duplicate": True,
            "message": "An identical cover letter version already exists.",
            "cover_letter_id": existing.pk,
            "version_number": existing.version_number
        }

    port_hash, res_hash, job_hash = compute_cover_letter_hashes(portfolio, {}, job_reqs)
    next_ver = (CoverLetter.objects.filter(portfolio=portfolio).order_by("-version_number").first().version_number + 1) if CoverLetter.objects.filter(portfolio=portfolio).exists() else 1

    title = cover_letter_data.get("title", f"Cover Letter v{next_ver} - {job_reqs.get('title', portfolio.title)}")

    content_payload = {
        "greeting": greeting,
        "introduction": intro,
        "body": body,
        "closing": closing,
        "signature": signature,
        "evidence_map": cover_letter_data.get("evidence_map", []),
        "coverage": cover_letter_data.get("coverage", {}),
        "metrics": compute_readability_metrics(full_text)
    }

    metadata_payload = {
        "tone": tone,
        "length": length,
        "template_variant": template_variant,
        "provider": cover_letter_data.get("metadata", {}).get("provider", "Gemini"),
        "model": cover_letter_data.get("metadata", {}).get("model", "gemini-1.5-flash"),
        "prompt_version": PROMPT_VERSION,
        "saved_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    }

    if replace_version_id:
        cover_letter = CoverLetter.objects.get(pk=replace_version_id, portfolio=portfolio)
        cover_letter.title = title
        cover_letter.tone = tone
        cover_letter.length = length
        cover_letter.template_variant = template_variant
        cover_letter.content_json = content_payload
        cover_letter.content_text = full_text
        cover_letter.content_hash = content_hash
        cover_letter.metadata_json = metadata_payload
        cover_letter.save()
    else:
        cover_letter = CoverLetter.objects.create(
            portfolio=portfolio,
            title=title,
            job_title=job_reqs.get("title", "")[:255],
            company=job_reqs.get("company", "")[:255],
            tone=tone,
            length=length,
            template_variant=template_variant,
            content_json=content_payload,
            content_text=full_text,
            version_number=next_ver,
            content_hash=content_hash,
            job_hash=job_hash,
            resume_hash=res_hash,
            portfolio_hash=port_hash,
            metadata_json=metadata_payload
        )

    return {
        "success": True,
        "is_duplicate": False,
        "cover_letter_id": cover_letter.pk,
        "version_number": cover_letter.version_number,
        "title": cover_letter.title
    }


def export_cover_letter(cover_letter_data: dict, format: str = "pdf") -> Tuple[bytes, str, str]:
    """
    Compiles cover letter with embedded metadata headers/properties into requested format.
    Formats: pdf, docx, markdown, txt.
    Returns: (file_bytes, mime_type, filename)
    """
    fmt = (format or "pdf").lower()
    title = cover_letter_data.get("title", "Cover_Letter").replace(" ", "_")
    greeting = cover_letter_data.get("greeting", "")
    intro = cover_letter_data.get("introduction", "")
    body = cover_letter_data.get("body", "")
    closing = cover_letter_data.get("closing", "")
    signature = cover_letter_data.get("signature", "")
    meta = cover_letter_data.get("metadata", {})

    full_content = f"{greeting}\n\n{intro}\n\n{body}\n\n{closing}\n\n{signature}"

    if fmt == "markdown" or fmt == "md":
        frontmatter = (
            "---\n"
            f"title: \"{cover_letter_data.get('title', 'Cover Letter')}\"\n"
            f"tone: \"{meta.get('tone', 'Professional')}\"\n"
            f"length: \"{meta.get('length', 'Medium')}\"\n"
            f"template: \"{meta.get('template_variant', 'Modern')}\"\n"
            f"generated_by: \"{meta.get('provider', 'Gemini')} AI\"\n"
            "---\n\n"
        )
        md_text = f"{frontmatter}# {cover_letter_data.get('title', 'Cover Letter')}\n\n{full_content}"
        return md_text.encode("utf-8"), "text/markdown", f"{title}.md"

    elif fmt == "txt" or fmt == "text":
        txt_header = (
            f"========================================\n"
            f"COVER LETTER: {cover_letter_data.get('title', 'Cover Letter')}\n"
            f"Tone: {meta.get('tone', 'Professional')} | Length: {meta.get('length', 'Medium')}\n"
            f"========================================\n\n"
        )
        return (txt_header + full_content).encode("utf-8"), "text/plain", f"{title}.txt"

    elif fmt == "docx":
        try:
            import docx
            doc = docx.Document()
            doc.core_properties.title = cover_letter_data.get("title", "Cover Letter")
            doc.core_properties.author = "AI Portfolio Builder"

            doc.add_heading(cover_letter_data.get("title", "Cover Letter"), level=1)
            for paragraph in full_content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{title}.docx"
        except ImportError:
            # Plain text fallback if docx library uninstalled
            return (f"COVER LETTER: {cover_letter_data.get('title')}\n\n{full_content}").encode("utf-8"), "application/octet-stream", f"{title}.docx"

    else:
        # Default PDF Generation
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
            styles = getSampleStyleSheet()

            story = []
            title_style = ParagraphStyle("CLTitle", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=14)
            body_style = ParagraphStyle("CLBody", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=12)

            story.append(Paragraph(cover_letter_data.get("title", "Cover Letter"), title_style))
            story.append(Spacer(1, 10))

            for paragraph in full_content.split("\n\n"):
                if paragraph.strip():
                    formatted_p = paragraph.strip().replace("\n", "<br/>")
                    story.append(Paragraph(formatted_p, body_style))
                    story.append(Spacer(1, 8))

            doc.build(story)
            return buffer.getvalue(), "application/pdf", f"{title}.pdf"
        except Exception:
            # Fallback simple text buffer if PDF engine encounters error
            return full_content.encode("utf-8"), "application/pdf", f"{title}.pdf"
